"""
==============================================================================
RAP NOVA ESCOLA — Serviço de Sincronização da Google Drive
==============================================================================
Descarrega todos os documentos suportados de uma pasta Drive e constrói
o ficheiro KNOWLEDGE_BASE.md usado pelo chatbot.

Formatos suportados:
  - PDF
  - Google Docs (exportado como texto)
  - DOCX / DOC
  - TXT / MD

Configuração necessária:
  - service_account.json na raiz do projeto
  - Variável DRIVE_FOLDER_ID (ou usa o valor fixo abaixo)
  - A pasta Drive deve estar partilhada com o email da service account
==============================================================================
"""

import io
import os
import logging
from datetime import datetime
from typing import List, Optional

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import pypdf
import docx

logger = logging.getLogger(__name__)

FOLDER_ID = os.getenv("DRIVE_FOLDER_ID", "1m6AMj2zO5Ce5w8ippfsEvuMyWBPOeABu")

_ROOT = os.path.dirname(os.path.dirname(__file__))
SERVICE_ACCOUNT_FILE = os.path.join(_ROOT, "service_account.json")
KNOWLEDGE_BASE_PATH = os.path.join(_ROOT, "KNOWLEDGE_BASE.md")

SUPPORTED_MIMES = {
    "application/pdf",
    "application/vnd.google-apps.document",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "text/markdown",
    "text/x-markdown",
}


# ---------------------------------------------------------------------------
# Drive client
# ---------------------------------------------------------------------------

def _get_drive_service():
    if not os.path.exists(SERVICE_ACCOUNT_FILE):
        raise FileNotFoundError(
            f"service_account.json não encontrado em {SERVICE_ACCOUNT_FILE}. "
            "Coloca o ficheiro JSON da Service Account na raiz do projeto."
        )

    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE,
        scopes=["https://www.googleapis.com/auth/drive.readonly"],
    )
    return build("drive", "v3", credentials=creds)


# ---------------------------------------------------------------------------
# Listagem recursiva de ficheiros
# ---------------------------------------------------------------------------

def _list_files(service, folder_id: str) -> List[dict]:
    """Lista todos os ficheiros suportados dentro da pasta (recursivo)."""
    results = []
    query = f"'{folder_id}' in parents and trashed = false"
    page_token = None

    while True:
        resp = service.files().list(
            q=query,
            fields="nextPageToken, files(id, name, mimeType)",
            pageToken=page_token,
            pageSize=200,
        ).execute()

        for f in resp.get("files", []):
            if f["mimeType"] == "application/vnd.google-apps.folder":
                results.extend(_list_files(service, f["id"]))
            elif f["mimeType"] in SUPPORTED_MIMES:
                results.append(f)

        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    return results


# ---------------------------------------------------------------------------
# Extração de texto por formato
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts = []

    # Parágrafos normais
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tabelas (linha a linha, células separadas por tab)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _download_and_extract(service, file_info: dict) -> Optional[str]:
    mime = file_info["mimeType"]
    name = file_info["name"]
    file_id = file_info["id"]

    try:
        if mime == "application/vnd.google-apps.document":
            # Google Docs: exportar como texto simples
            raw = service.files().export(fileId=file_id, mimeType="text/plain").execute()
            text = raw.decode("utf-8") if isinstance(raw, bytes) else str(raw)

        else:
            # Todos os outros: download binário
            request = service.files().get_media(fileId=file_id)
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw_bytes = buf.getvalue()

            if mime == "application/pdf":
                text = _extract_pdf(raw_bytes)
            elif mime in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ):
                text = _extract_docx(raw_bytes)
            else:
                text = raw_bytes.decode("utf-8", errors="replace")

        text = text.strip()
        if not text:
            return None

        return f'# FONTE: "{name}"\n\n{text}'

    except Exception as exc:
        logger.warning(f"Não foi possível processar '{name}': {exc}")
        return None


# ---------------------------------------------------------------------------
# Função principal de sync
# ---------------------------------------------------------------------------

def sync_knowledge_base() -> dict:
    """
    Descarrega todos os documentos da pasta Drive e reconstrói o KNOWLEDGE_BASE.md.
    Retorna um dict com estatísticas do sync.
    """
    logger.info("Drive sync → KNOWLEDGE_BASE: a iniciar...")

    service = _get_drive_service()
    files = _list_files(service, FOLDER_ID)
    logger.info(f"Ficheiros encontrados na Drive: {len(files)}")

    sections: List[str] = []
    errors = 0

    for f in files:
        text = _download_and_extract(service, f)
        if text:
            sections.append(text)
            logger.debug(f"  ✓ {f['name']}")
        else:
            errors += 1
            logger.debug(f"  ✗ {f['name']} (ignorado)")

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    header = f"<!-- Base de Conhecimento — Sincronizada em {timestamp} -->\n\n"
    combined = header + "\n\n---\n\n".join(sections)

    with open(KNOWLEDGE_BASE_PATH, "w", encoding="utf-8") as fp:
        fp.write(combined)

    stats = {
        "synced_files": len(sections),
        "skipped_files": errors,
        "total_files": len(files),
        "kb_chars": len(combined),
        "synced_at": datetime.now().isoformat(),
    }
    logger.info(f"Drive sync concluído: {stats}")
    return stats
